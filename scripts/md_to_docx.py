"""Convert REPORT.md → REPORT.docx using python-docx.

Minimal Markdown subset: ATX headings (#..######), unordered lists (- ),
ordered lists (1. ), inline code `x`, bold **x**, italics *x*, fenced code
blocks ```...```, GitHub-style tables (| a | b | with --- separator), blank
lines as paragraph breaks. Good enough for our REPORT.md.
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
# Allow override: `python scripts/md_to_docx.py SUBMISSION_REPORT.md`
SRC = ROOT / (sys.argv[1] if len(sys.argv) > 1 else "REPORT.md")
DST = SRC.with_suffix(".docx")


INLINE_CODE = re.compile(r"`([^`]+)`")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
ITALIC = re.compile(r"(?<!\*)\*([^*]+)\*(?!\*)")
LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def add_runs(paragraph, text: str) -> None:
    """Render bold / italics / inline code / links inside a paragraph."""
    text = LINK.sub(r"\1 (\2)", text)
    # Tokenize on the three markers, preserving order.
    tokens: list[tuple[str, str]] = []
    pos = 0
    pattern = re.compile(r"(\*\*[^*]+\*\*|(?<!\*)\*[^*]+\*(?!\*)|`[^`]+`)")
    for m in pattern.finditer(text):
        if m.start() > pos:
            tokens.append(("plain", text[pos:m.start()]))
        chunk = m.group(0)
        if chunk.startswith("**"):
            tokens.append(("bold", chunk[2:-2]))
        elif chunk.startswith("`"):
            tokens.append(("code", chunk[1:-1]))
        else:
            tokens.append(("italic", chunk[1:-1]))
        pos = m.end()
    if pos < len(text):
        tokens.append(("plain", text[pos:]))

    for kind, val in tokens:
        run = paragraph.add_run(val)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True
        elif kind == "code":
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x33, 0x66, 0xCC)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Read a Markdown table starting at `start`. Returns (rows, next_index)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        rows.append(cells)
        i += 1
    # Drop the separator row (---|---).
    rows = [r for r in rows if not all(re.fullmatch(r":?-+:?", c) for c in r)]
    return rows, i


def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_buf))
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Figure marker:  <!-- FIG: path | Caption -->
        m_fig = re.match(r"^<!--\s*FIG:\s*([^|]+?)\s*\|\s*(.+?)\s*-->$", stripped)
        if m_fig:
            rel, caption = m_fig.group(1).strip(), m_fig.group(2).strip()
            img_path = (ROOT / rel).resolve()
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(str(img_path), width=Inches(6.4))
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.italic = True
                run.font.size = Pt(9)
            else:
                doc.add_paragraph(f"[Missing figure: {rel}]").italic = True
            i += 1
            continue

        # Horizontal rule
        if stripped in {"---", "***"}:
            doc.add_paragraph().add_run("─" * 60)
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            add_runs(p, stripped.lstrip("> ").rstrip())
            i += 1
            continue

        # Table
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|?\s*:?-+:?", lines[i + 1]
        ):
            rows, next_i = parse_table(lines, i)
            if rows:
                tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
                tbl.style = "Light Grid Accent 1"
                for r, row in enumerate(rows):
                    for c, val in enumerate(row):
                        cell = tbl.rows[r].cells[c]
                        cell.text = ""
                        para = cell.paragraphs[0]
                        add_runs(para, val)
                        if r == 0:
                            for run in para.runs:
                                run.bold = True
            i = next_i
            continue

        # Unordered list
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Paragraph (consume continuation lines until blank/structural)
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt or re.match(
                r"^(#{1,6}\s|[-*]\s|\d+\.\s|\||```|>)", nxt
            ):
                break
            buf.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(buf))

    doc.save(DST)
    print(f"Wrote {DST} ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    sys.exit(main() or 0)
